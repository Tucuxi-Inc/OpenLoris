"""
Document Service — handles document upload, parsing, chunking, fact extraction,
and candidate approval/rejection workflow.

Ported from CounselScope's document_ingestion_service.py,
adapted to Loris async patterns (AsyncSession passed in, UUIDs, Mapped models).
"""

import logging
import os
import re
import uuid as _uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from uuid import UUID

from sqlalchemy import and_, delete, desc, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.documents import (
    ChunkEmbedding,
    DocumentChunk,
    DocumentType,
    ExtractionStatus,
    ExtractedFactCandidate,
    KnowledgeDocument,
    ParsingStatus,
    ValidationStatus,
)
from app.models.wisdom import WisdomFact, WisdomTier
from app.services.embedding_service import embedding_service

logger = logging.getLogger(__name__)

# Try optional parsing dependencies
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logger.warning("pdfplumber not installed — PDF parsing disabled")

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("python-docx not installed — DOCX parsing disabled")


class DocumentService:
    """Ingestion, parsing, fact extraction, and management of knowledge documents."""

    def __init__(self):
        upload_dir = getattr(settings, "UPLOAD_DIR", "/app/uploads")
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

        self.max_chunk_size = 2000  # characters
        self.chunk_overlap = 200

    # ------------------------------------------------------------------
    # Upload / ingest
    # ------------------------------------------------------------------

    async def ingest_document(
        self,
        db: AsyncSession,
        file_bytes: bytes,
        original_filename: str,
        organization_id: UUID,
        uploaded_by_id: UUID,
        *,
        document_type: str = "other",
        domain: Optional[str] = None,
        title: Optional[str] = None,
        description: Optional[str] = None,
        topics: Optional[list] = None,
        tags: Optional[list] = None,
        department: Optional[str] = None,
        responsible_person: Optional[str] = None,
        responsible_email: Optional[str] = None,
        good_until_date: Optional[str] = None,
        is_perpetual: bool = True,
        auto_delete_on_expiry: bool = False,
    ) -> Tuple[Optional[KnowledgeDocument], Optional[str]]:
        """
        Save file, create KnowledgeDocument record, parse, and chunk.
        Returns (document, error_message).
        """
        try:
            file_ext = Path(original_filename).suffix.lower().lstrip(".")
            stored_name = f"{_uuid.uuid4().hex}.{file_ext}"
            stored_path = self.upload_dir / stored_name

            # Write file
            with open(stored_path, "wb") as f:
                f.write(file_bytes)

            # Determine GUD date
            gud = None
            if good_until_date:
                from datetime import date as _date
                try:
                    gud = _date.fromisoformat(good_until_date)
                except ValueError:
                    pass

            doc = KnowledgeDocument(
                organization_id=organization_id,
                uploaded_by_id=uploaded_by_id,
                original_filename=original_filename,
                file_path=str(stored_path),
                file_size_bytes=len(file_bytes),
                file_type=file_ext,
                mime_type=self._guess_mime(file_ext),
                document_type=DocumentType(document_type) if document_type in [e.value for e in DocumentType] else DocumentType.OTHER,
                title=title or original_filename,
                description=description,
                domain=domain,
                topics=topics,
                tags=tags,
                department=department,
                responsible_person=responsible_person,
                responsible_email=responsible_email,
                good_until_date=gud,
                is_perpetual=is_perpetual,
                auto_delete_on_expiry=auto_delete_on_expiry,
                parsing_status=ParsingStatus.PROCESSING,
                parsing_started_at=datetime.now(timezone.utc),
                extraction_status=ExtractionStatus.PENDING,
            )
            db.add(doc)
            await db.flush()

            # Parse
            content, parse_error = await self._parse_document(str(stored_path), file_ext)

            if parse_error or not content:
                doc.parsing_status = ParsingStatus.FAILED
                doc.parsing_error = parse_error or "No content extracted"
                await db.commit()
                return doc, doc.parsing_error

            doc.word_count = len(content.split())

            # Chunk
            chunks = self._create_chunks(content)
            doc.chunk_count = len(chunks)

            for idx, chunk_text in enumerate(chunks):
                chunk = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=idx,
                    content=chunk_text,
                    word_count=len(chunk_text.split()),
                    content_type="text",
                )
                db.add(chunk)

            doc.parsing_status = ParsingStatus.COMPLETED
            doc.parsing_completed_at = datetime.now(timezone.utc)

            await db.commit()
            await db.refresh(doc)

            logger.info(f"Ingested document {doc.id} ({original_filename}), {len(chunks)} chunks")
            return doc, None

        except Exception as e:
            logger.error(f"Error ingesting document: {e}")
            await db.rollback()
            return None, str(e)

    # ------------------------------------------------------------------
    # Parse helpers
    # ------------------------------------------------------------------

    async def _parse_document(self, file_path: str, file_type: str) -> Tuple[Optional[str], Optional[str]]:
        try:
            if file_type == "pdf":
                return await self._parse_pdf(file_path)
            elif file_type in ("docx", "doc"):
                return await self._parse_docx(file_path)
            elif file_type in ("txt", "md", "markdown"):
                return await self._parse_text(file_path)
            else:
                return None, f"Unsupported file type: {file_type}"
        except Exception as e:
            return None, str(e)

    async def _parse_pdf(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        if not PDF_SUPPORT:
            return None, "pdfplumber not installed"
        try:
            pages = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n\n".join(pages), None
        except Exception as e:
            return None, f"PDF parsing error: {e}"

    async def _parse_docx(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        if not DOCX_SUPPORT:
            return None, "python-docx not installed"
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs), None
        except Exception as e:
            return None, f"DOCX parsing error: {e}"

    async def _parse_text(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read(), None
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    return f.read(), None
            except Exception as e:
                return None, f"Text parsing error: {e}"
        except Exception as e:
            return None, f"Text parsing error: {e}"

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def _create_chunks(self, content: str) -> List[str]:
        if len(content) <= self.max_chunk_size:
            return [content]

        chunks: List[str] = []
        paragraphs = content.split("\n\n")
        current = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) + 2 > self.max_chunk_size:
                if current:
                    chunks.append(current)
                    overlap = current[-self.chunk_overlap:] if len(current) > self.chunk_overlap else current
                    current = overlap + "\n\n" + para
                else:
                    # Single paragraph exceeds max
                    words = para.split()
                    temp = ""
                    for word in words:
                        if len(temp) + len(word) + 1 > self.max_chunk_size:
                            chunks.append(temp)
                            temp = word
                        else:
                            temp = f"{temp} {word}" if temp else word
                    current = temp
            else:
                current = f"{current}\n\n{para}" if current else para

        if current:
            chunks.append(current)
        return chunks

    # ------------------------------------------------------------------
    # Fact extraction
    # ------------------------------------------------------------------

    async def extract_facts(
        self,
        db: AsyncSession,
        document_id: UUID,
    ) -> Tuple[int, Optional[str]]:
        """
        Extract fact candidates from a parsed document using rule-based patterns.
        Returns (count_extracted, error_message).
        """
        result = await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
        )
        doc = result.scalar_one_or_none()
        if not doc:
            return 0, "Document not found"
        if doc.parsing_status != ParsingStatus.COMPLETED:
            return 0, "Document not fully parsed"

        doc.extraction_status = ExtractionStatus.EXTRACTING
        await db.flush()

        chunks_result = await db.execute(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.chunk_index)
        )
        chunks = chunks_result.scalars().all()

        total = 0
        for chunk in chunks:
            facts = self._extract_facts_from_text(chunk.content)
            for fact_data in facts:
                candidate = ExtractedFactCandidate(
                    document_id=document_id,
                    chunk_id=chunk.id,
                    fact_text=fact_data["fact_text"],
                    summary=fact_data.get("summary"),
                    source_excerpt=fact_data.get("source_excerpt"),
                    suggested_domain=doc.domain,
                    suggested_importance=fact_data.get("importance", 5),
                    suggested_tags=fact_data.get("tags"),
                    extraction_confidence=fact_data.get("confidence", 0.5),
                    validation_status=ValidationStatus.PENDING,
                )
                db.add(candidate)
                total += 1
            chunk.is_processed = True
            chunk.facts_extracted = len(facts)

        doc.extracted_facts_count = total
        doc.extraction_status = ExtractionStatus.COMPLETED if total > 0 else ExtractionStatus.PARTIAL

        await db.commit()
        logger.info(f"Extracted {total} fact candidates from document {document_id}")
        return total, None

    def _extract_facts_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Rule-based fact extraction (production would use LLM)."""
        facts: List[Dict[str, Any]] = []
        sentences = re.split(r"(?<=[.!?])\s+", text)

        patterns = [
            (r"^(?:our|the|company)\s+policy\s+(?:is|requires|states)", 9),
            (r"(?:must|shall|required\s+to|need\s+to)\s+", 8),
            (r"(?:compliance|regulatory|legal\s+requirement|pursuant\s+to)", 8),
            (r"(?:contract|agreement|terms)\s+(?:should|must|require)", 7),
            (r"(?:threshold|limit|maximum|minimum|cap)\s+(?:of|is|at)", 7),
            (r"(?:approval|sign-off|authorization)\s+(?:required|needed|from)", 7),
            (r"(?:standard\s+practice|typically|generally|always)\s+", 5),
        ]

        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 30 or len(sentence) > 500:
                continue
            for pattern, importance in patterns:
                if re.search(pattern, sentence, re.IGNORECASE):
                    cleaned = self._clean_fact(sentence)
                    if cleaned and len(cleaned) > 30:
                        tags = self._extract_tags(cleaned)
                        facts.append({
                            "fact_text": cleaned,
                            "summary": cleaned[:200] + "..." if len(cleaned) > 200 else cleaned,
                            "source_excerpt": sentence,
                            "importance": importance,
                            "confidence": 0.6,
                            "tags": tags,
                        })
                    break
        return facts

    @staticmethod
    def _clean_fact(text: str) -> str:
        text = " ".join(text.split())
        text = re.sub(r"^[\s•\-*\d.]+", "", text)
        if text and text[-1] not in ".!?":
            text += "."
        return text.strip()

    @staticmethod
    def _extract_tags(text: str) -> List[str]:
        lower = text.lower()
        tag_terms = {
            "contract": ["contract", "agreement", "terms"],
            "compliance": ["compliance", "regulatory", "requirement"],
            "privacy": ["privacy", "data protection", "gdpr", "ccpa"],
            "liability": ["liability", "indemnification", "limitation"],
            "ip": ["intellectual property", "patent", "trademark", "copyright"],
            "employment": ["employment", "employee", "labor"],
            "security": ["security", "breach", "incident"],
            "approval": ["approval", "authorization", "sign-off"],
        }
        tags = []
        for tag, terms in tag_terms.items():
            if any(t in lower for t in terms):
                tags.append(tag)
        return tags[:5]

    # ------------------------------------------------------------------
    # Candidate approval / rejection
    # ------------------------------------------------------------------

    async def approve_candidate(
        self,
        db: AsyncSession,
        candidate_id: UUID,
        expert_user_id: UUID,
        organization_id: UUID,
        *,
        modified_text: Optional[str] = None,
        domain: Optional[str] = None,
        importance: Optional[int] = None,
    ) -> Tuple[Optional[WisdomFact], Optional[str]]:
        """
        Approve a fact candidate → create WisdomFact + embedding.
        Returns (wisdom_fact, error).
        """
        from app.services.knowledge_service import knowledge_service

        result = await db.execute(
            select(ExtractedFactCandidate).where(ExtractedFactCandidate.id == candidate_id)
        )
        candidate = result.scalar_one_or_none()
        if not candidate:
            return None, "Candidate not found"
        if candidate.validation_status != ValidationStatus.PENDING:
            return None, f"Candidate already {candidate.validation_status.value}"

        fact = await knowledge_service.create_fact(
            db=db,
            organization_id=organization_id,
            content=modified_text or candidate.fact_text,
            expert_user_id=expert_user_id,
            summary=candidate.summary,
            domain=domain or candidate.suggested_domain,
            importance=importance or candidate.suggested_importance or 5,
            tags=candidate.suggested_tags,
            confidence_score=candidate.extraction_confidence,
            tier=WisdomTier.TIER_0C,
            source_document_id=candidate.document_id,
        )

        candidate.validation_status = ValidationStatus.APPROVED
        candidate.validated_by_id = expert_user_id
        candidate.validated_at = datetime.now(timezone.utc)
        candidate.created_wisdom_fact_id = fact.id

        # Update document counters
        await db.execute(
            select(KnowledgeDocument)
            .where(KnowledgeDocument.id == candidate.document_id)
        )
        doc_result = await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == candidate.document_id)
        )
        doc = doc_result.scalar_one_or_none()
        if doc:
            doc.validated_facts_count = (doc.validated_facts_count or 0) + 1

        await db.commit()
        return fact, None

    async def reject_candidate(
        self,
        db: AsyncSession,
        candidate_id: UUID,
        expert_user_id: UUID,
        reason: str,
    ) -> bool:
        result = await db.execute(
            select(ExtractedFactCandidate).where(ExtractedFactCandidate.id == candidate_id)
        )
        candidate = result.scalar_one_or_none()
        if not candidate:
            return False

        candidate.validation_status = ValidationStatus.REJECTED
        candidate.validated_by_id = expert_user_id
        candidate.validated_at = datetime.now(timezone.utc)
        candidate.rejection_reason = reason
        await db.commit()
        return True

    # ------------------------------------------------------------------
    # Queries
    # ------------------------------------------------------------------

    async def get_document(
        self, db: AsyncSession, document_id: UUID,
    ) -> Optional[KnowledgeDocument]:
        result = await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
        )
        return result.scalar_one_or_none()

    async def list_documents(
        self,
        db: AsyncSession,
        organization_id: UUID,
        *,
        status: Optional[str] = None,
        page: int = 1,
        page_size: int = 20,
    ) -> Dict[str, Any]:
        stmt = select(KnowledgeDocument).where(
            KnowledgeDocument.organization_id == organization_id,
            KnowledgeDocument.is_active == True,
        )
        if status:
            try:
                stmt = stmt.where(KnowledgeDocument.parsing_status == ParsingStatus(status))
            except ValueError:
                pass

        count_stmt = select(func.count()).select_from(stmt.subquery())
        total = (await db.execute(count_stmt)).scalar() or 0

        offset = (page - 1) * page_size
        stmt = stmt.order_by(desc(KnowledgeDocument.created_at)).offset(offset).limit(page_size)
        rows = (await db.execute(stmt)).scalars().all()

        return {
            "documents": [self._doc_to_dict(d) for d in rows],
            "total": total,
            "page": page,
            "page_size": page_size,
            "total_pages": max(1, (total + page_size - 1) // page_size),
        }

    async def get_candidates(
        self,
        db: AsyncSession,
        document_id: UUID,
        status: Optional[str] = None,
    ) -> List[Dict[str, Any]]:
        stmt = select(ExtractedFactCandidate).where(
            ExtractedFactCandidate.document_id == document_id
        )
        if status:
            try:
                stmt = stmt.where(ExtractedFactCandidate.validation_status == ValidationStatus(status))
            except ValueError:
                pass
        stmt = stmt.order_by(desc(ExtractedFactCandidate.extraction_confidence))

        rows = (await db.execute(stmt)).scalars().all()
        return [
            {
                "id": str(c.id),
                "document_id": str(c.document_id),
                "fact_text": c.fact_text,
                "summary": c.summary,
                "source_excerpt": c.source_excerpt,
                "suggested_domain": c.suggested_domain,
                "suggested_importance": c.suggested_importance,
                "suggested_tags": c.suggested_tags,
                "extraction_confidence": c.extraction_confidence,
                "validation_status": c.validation_status.value if c.validation_status else "pending",
                "rejection_reason": c.rejection_reason,
                "created_at": c.created_at.isoformat() if c.created_at else None,
            }
            for c in rows
        ]

    async def delete_document(
        self,
        db: AsyncSession,
        document_id: UUID,
    ) -> bool:
        result = await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
        )
        doc = result.scalar_one_or_none()
        if not doc:
            return False

        # Remove file
        if doc.file_path and os.path.exists(doc.file_path):
            try:
                os.remove(doc.file_path)
            except Exception as e:
                logger.warning(f"Could not delete file {doc.file_path}: {e}")

        await db.delete(doc)
        await db.commit()
        return True

    async def update_document(
        self,
        db: AsyncSession,
        document_id: UUID,
        updates: Dict[str, Any],
    ) -> Optional[KnowledgeDocument]:
        result = await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
        )
        doc = result.scalar_one_or_none()
        if not doc:
            return None

        allowed = {
            "title", "description", "domain", "topics", "tags",
            "department", "responsible_person", "responsible_email",
            "good_until_date", "is_perpetual", "auto_delete_on_expiry",
            "document_type",
        }
        for key, val in updates.items():
            if key in allowed and hasattr(doc, key):
                if key == "document_type":
                    try:
                        val = DocumentType(val)
                    except ValueError:
                        continue
                if key == "good_until_date" and isinstance(val, str):
                    from datetime import date as _date
                    try:
                        val = _date.fromisoformat(val)
                    except ValueError:
                        continue
                setattr(doc, key, val)

        await db.commit()
        await db.refresh(doc)
        return doc

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _doc_to_dict(doc: KnowledgeDocument) -> Dict[str, Any]:
        return {
            "id": str(doc.id),
            "original_filename": doc.original_filename,
            "title": doc.title,
            "description": doc.description,
            "document_type": doc.document_type.value if doc.document_type else "other",
            "domain": doc.domain,
            "parsing_status": doc.parsing_status.value if doc.parsing_status else None,
            "extraction_status": doc.extraction_status.value if doc.extraction_status else None,
            "chunk_count": doc.chunk_count,
            "word_count": doc.word_count,
            "extracted_facts_count": doc.extracted_facts_count,
            "validated_facts_count": doc.validated_facts_count,
            "department": doc.department,
            "responsible_person": doc.responsible_person,
            "good_until_date": doc.good_until_date.isoformat() if doc.good_until_date else None,
            "is_perpetual": doc.is_perpetual,
            "auto_delete_on_expiry": doc.auto_delete_on_expiry,
            "is_active": doc.is_active,
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
        }

    @staticmethod
    def _guess_mime(ext: str) -> str:
        return {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "doc": "application/msword",
            "txt": "text/plain",
            "md": "text/markdown",
        }.get(ext, "application/octet-stream")


# Global singleton
document_service = DocumentService()
